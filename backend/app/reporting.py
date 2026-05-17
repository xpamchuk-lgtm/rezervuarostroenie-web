from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .engine.rgs_engine import calculate_rgs
from .engine.strength_engine import calculate_strength

ACCENT = "1F4E79"
ACCENT_LIGHT = "D9E8F5"
BORDER = "9FBBD3"
RED_LIGHT = "FDECEC"
YELLOW_LIGHT = "FFF4CE"
GREEN_LIGHT = "E7F4EA"


def _safe_filename(prefix: str) -> str:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefix}_{stamp}.docx"


def _fmt(value: float | int | None, digits: int = 3) -> str:
    if value is None:
        return "—"
    return f"{float(value):.{digits}f}".replace(".", ",")


def _fmt_any(value: Any, digits: int = 3) -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, bool):
        return "да" if value else "нет"
    if isinstance(value, (int, float)):
        return _fmt(value, digits)
    return str(value)


def _bool_text(value: bool) -> str:
    return "соответствует" if value else "не соответствует"


def _status_text(value: str) -> str:
    return {
        "PASS": "OK",
        "FAIL": "НЕ ОК",
        "critical": "Критично",
        "important": "Важно",
        "advice": "Рекомендация",
    }.get(value, value)


class DocxBuilder:
    def __init__(self, doc: Document):
        self.doc = doc
        self._setup()

    def _setup(self) -> None:
        sec = self.doc.sections[0]
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(1.8)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10.5)

    def _set_para_spacing(self, p, before=0, after=0, line=1.08):
        p_format = p.paragraph_format
        p_format.space_before = Pt(before)
        p_format.space_after = Pt(after)
        p_format.line_spacing = line
        return p

    def para(self, text: str = "", *, bold=False, italic=False, center=False, color: str | None = None, size: float = 10.5):
        p = self.doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        self._set_para_spacing(p, after=3)
        return p

    def bullet(self, text: str):
        p = self.doc.add_paragraph(style=None)
        p.style = self.doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        run = p.add_run(f"• {text}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        self._set_para_spacing(p, after=2)
        return p

    def heading(self, text: str, level: int = 1):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14 if level == 1 else 11.5)
        run.font.color.rgb = RGBColor.from_string(ACCENT)
        self._set_para_spacing(p, before=6 if level == 1 else 4, after=4)
        return p

    def page_break(self):
        self.doc.add_page_break()

    def title_box(self, title: str, subtitle: str, meta_rows: Sequence[tuple[str, str]]):
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Cm(16.5)
        self._shade(cell, ACCENT)
        self._set_cell_border(cell, color=ACCENT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor(255, 255, 255)
        self._set_para_spacing(p, after=2)
        self._set_para_spacing(p2, after=0)

        self.para("")
        self.key_value_table(meta_rows, first_col_width_cm=6.0, second_col_width_cm=10.0)

    def key_value_table(self, rows: Sequence[tuple[str, str]], *, title: str | None = None, first_col_width_cm: float = 7.0, second_col_width_cm: float = 9.0):
        if title:
            self.heading(title, 1)
        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_layout(table)
        hdr = table.rows[0].cells
        self._set_cell(hdr[0], "Показатель", bold=True, fill=ACCENT_LIGHT, width_cm=first_col_width_cm)
        self._set_cell(hdr[1], "Значение", bold=True, fill=ACCENT_LIGHT, width_cm=second_col_width_cm)
        for key, value in rows:
            row = table.add_row().cells
            self._set_cell(row[0], key, width_cm=first_col_width_cm)
            self._set_cell(row[1], value, width_cm=second_col_width_cm)
        self.para("")
        return table

    def grid_table(self, headers: Sequence[str], rows: Sequence[Sequence[str]], *, title: str | None = None, widths_cm: Sequence[float] | None = None, status_col: int | None = None):
        if title:
            self.heading(title, 1)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_layout(table)
        for idx, text in enumerate(headers):
            self._set_cell(table.rows[0].cells[idx], text, bold=True, fill=ACCENT_LIGHT, width_cm=widths_cm[idx] if widths_cm else None, center=True)
        for row_values in rows:
            row = table.add_row().cells
            for idx, value in enumerate(row_values):
                fill = None
                if status_col is not None and idx == status_col:
                    value_n = value.lower()
                    if "крит" in value_n or "не ок" in value_n or "fail" in value_n:
                        fill = RED_LIGHT
                    elif "важ" in value_n:
                        fill = YELLOW_LIGHT
                    elif "ok" in value_n or "соответ" in value_n or "рекомендац" in value_n:
                        fill = GREEN_LIGHT if "ok" in value_n or "соответ" in value_n else ACCENT_LIGHT
                self._set_cell(row[idx], value, width_cm=widths_cm[idx] if widths_cm else None, fill=fill, center=(idx == status_col or (len(value) <= 12 and idx != 1)))
        self.para("")
        return table

    def note_box(self, title: str, text: str, fill: str = YELLOW_LIGHT):
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Cm(16.5)
        self._shade(cell, fill)
        self._set_cell_border(cell, color=BORDER)
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{title}: ")
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10.5)
        self._set_para_spacing(p, after=0)
        self.para("")
        return table

    def _set_table_layout(self, table):
        tbl_pr = table._tbl.tblPr
        tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

    def _shade(self, cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.first_child_found_in("w:shd")
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def _set_cell_border(self, cell, *, color: str = BORDER):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            el = tc_borders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tc_borders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)

    def _set_cell(self, cell, text: str, *, bold=False, fill: str | None = None, width_cm: float | None = None, center=False):
        if fill:
            self._shade(cell, fill)
        self._set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if width_cm is not None:
            cell.width = Cm(width_cm)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.clear()
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)
        self._set_para_spacing(p, after=0, line=1.0)


def _build_strength_report(doc: Document, result: dict) -> None:
    b = DocxBuilder(doc)
    meta = result["meta"]
    inputs = result["inputs"]
    geometry = result["geometry"]
    bottom = result["bottom"]
    roof = result["roof"]
    masses = result["masses"]
    stability = result["stability"]
    summary = result["summary"]
    completeness = result["data_completeness"]
    design_basis = result["design_basis"]
    flags = result["review_flags"]

    b.title_box(
        "ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ ПО РАСЧЕТНОЙ МОДЕЛИ РВС",
        result["notes"]["report_type"],
        [
            ("Объект", meta.get("title") or "РВС"),
            ("Тип резервуара", meta.get("tank_type") or "РВС"),
            ("Класс резервуара", meta.get("reservoir_class") or "—"),
            ("Геометрия", f"D = {_fmt(geometry['diameter_m'], 3)} м; H = {_fmt(geometry['height_m'], 3)} м"),
            ("Среда", inputs.get("medium") or "—"),
            ("Дата формирования", datetime.now().strftime("%d.%m.%Y %H:%M:%S")),
        ],
    )

    b.note_box(
        "Статус отчета",
        "Документ оформлен как самостоятельная экспертная проверка расчетной модели. Его структура намеренно отличается от типовой пояснительной записки ПАССАТ и не воспроизводит ее оформление или последовательность разделов.",
        fill=ACCENT_LIGHT,
    )

    b.heading("1. Примененная нормативная база", 1)
    for item in result["normative"]:
        b.bullet(item)
    b.note_box(
        "Важно",
        "При наличии специальных условий эксплуатации следует применять профильные документы по промышленной безопасности, пожарной защите, КИПиА, технологической части и инженерным изысканиям сверх данного расчета.",
    )

    b.heading("2. Границы применимости и объем автоматизированной проверки", 1)
    b.para(result["notes"]["calculation_scope"])
    b.heading("2.1 Что в отчете действительно проверено", 2)
    for item in result["scope"]["performed"]:
        b.bullet(item)
    b.heading("2.2 Что остается вне автоматизированного ядра и должно быть закрыто проектом", 2)
    for item in result["scope"]["pending"]:
        b.bullet(item)

    b.key_value_table([
        ("Номинальный объем, м³", _fmt(geometry["full_volume_m3"], 3)),
        ("Полезный объем, м³", _fmt(geometry["useful_volume_m3"], 3)),
        ("Расчетный уровень налива, мм", _fmt(inputs["fill_level_mm"], 0)),
        ("Плотность продукта, кг/м³", _fmt(inputs["density_kg_m3"], 1)),
        ("Избыточное давление, МПа", _fmt(inputs["p_gas_mpa"], 5)),
        ("Испытательное давление, МПа", _fmt(inputs["p_test_mpa"], 5)),
        ("Относительный вакуум, кПа", _fmt(inputs.get("vacuum_kpa"), 3)),
        ("Ветровое давление, кПа", _fmt(inputs["wind_kpa"], 3)),
        ("Снеговая нагрузка, кПа", _fmt(inputs["snow_kpa"], 3)),
        ("Сейсмичность", str(inputs["seismic"])),
        ("Допускаемое давление на основание, кПа", _fmt(inputs.get("foundation_limit_kpa"), 3)),
        ("Коэффициент трения", _fmt(inputs.get("friction_coeff"), 3)),
        ("ГО/УЛФ / инертирование", str(inputs.get("gas_balance_system") or "—")),
        ("Срок службы, лет", _fmt(inputs.get("service_life_years"), 0) if inputs.get("service_life_years") else "—"),
        ("Циклов в год", _fmt(inputs.get("cycles_per_year"), 0) if inputs.get("cycles_per_year") else "—"),
    ], title="3. Исходная расчетная база")

    completeness_rows = []
    for item in completeness["items"]:
        completeness_rows.append([
            item["title"],
            "заполнено" if item["filled"] else "не заполнено",
            "критично" if item["critical"] else "дополнительно",
            item["value"] or "—",
        ])
    b.grid_table(
        ["Параметр", "Статус", "Приоритет", "Значение / комментарий"],
        completeness_rows,
        title=f"4. Полнота исходных данных по перечню ГОСТ 31385-2023 (заполнено {completeness['filled']} из {completeness['total']}, {completeness['percent']} %)",
        widths_cm=[6.0, 2.8, 2.6, 6.1],
        status_col=1,
    )

    if completeness["missing"]:
        b.note_box(
            "Замечание по данным",
            "До статуса полноценной проектной записки необходимо закрыть отсутствующие обязательные поля задания на проектирование и геотехнической части.",
        )

    belt_rows = []
    for belt in result["belts"]:
        belt_rows.append([
            str(belt["belt"]),
            _fmt(belt["height_mm"], 0),
            _fmt(belt["thickness_nominal_mm"], 2),
            _fmt(belt["thickness_effective_mm"], 2),
            _fmt(belt["thickness_required_oper_mm"], 2),
            _fmt(belt["thickness_required_test_mm"], 2),
            _fmt(belt["thickness_required_nominal_mm"], 2),
            _fmt(belt["reserve_ratio"], 3),
            _bool_text(bool(belt["ok"])),
        ])
    b.grid_table(
        ["Пояс", "h, мм", "tном, мм", "tэф, мм", "tтр экспл, мм", "tтр гидр, мм", "tтр ном, мм", "Запас", "Статус"],
        belt_rows,
        title="5. Проверка поясов стенки",
        widths_cm=[1.1, 1.5, 1.6, 1.6, 2.2, 2.2, 2.2, 1.6, 1.9],
        status_col=8,
    )
    b.para(
        "Принцип оценки: для каждого пояса определена требуемая номинальная толщина с учетом минимальной конструктивной толщины по ГОСТ 31385-2023, припуска на коррозию и минусового допуска на прокат. В таблицу вынесены отдельно эксплуатационный и гидроиспытательный случаи.",
        italic=True,
    )

    b.key_value_table([
        ("Исполнение днища", str(bottom["execution"])),
        ("Диаметр днища, мм", _fmt(bottom["bottom_diameter_mm"], 0)),
        ("Ширина окрайки, мм", _fmt(bottom["ring_width_mm"], 0)),
        ("Толщина центральной части, мм", _fmt(bottom["bottom_t_mm"], 2)),
        ("Толщина окрайки, мм", _fmt(bottom["ring_t_mm"], 2)),
        ("Припуск на коррозию днища, мм", _fmt(bottom.get("bottom_corr_mm"), 2)),
        ("Масса днища, кг", _fmt(bottom["total_mass_kg"], 1)),
    ], title="6. Днище")

    b.key_value_table([
        ("Тип кровли", str(roof["type"])),
        ("Угол наклона, град", _fmt(roof["angle_deg"], 2)),
        ("Толщина настила, мм", _fmt(roof["deck_t_mm"], 2)),
        ("Припуск на коррозию крыши, мм", _fmt(roof.get("deck_corr_mm"), 2)),
        ("Масса кровли, кг", _fmt(roof["mass_kg"], 1)),
        ("Снеговая нагрузка, кПа", _fmt(inputs["snow_kpa"], 3)),
    ], title="7. Крыша")

    b.key_value_table([
        ("Масса стенки, кг", _fmt(masses["shell_kg"], 1)),
        ("Масса днища, кг", _fmt(masses["bottom_kg"], 1)),
        ("Масса кровли, кг", _fmt(masses["roof_kg"], 1)),
        ("Масса металлоконструкций, кг", _fmt(masses["metal_kg"], 1)),
        ("Масса теплоизоляции, кг", _fmt(masses["insulation_kg"], 1)),
        ("Масса продукта, кг", _fmt(masses["product_kg"], 1)),
        ("Масса снега, кг", _fmt(masses["snow_kg"], 1)),
        ("Суммарная масса, кг", _fmt(masses["total_kg"], 1)),
        ("Суммарная вертикальная нагрузка, кН", _fmt(masses["total_load_kn"], 3)),
    ], title="8. Массы и вертикальные нагрузки")

    b.key_value_table([
        ("Ветровая сила, кН", _fmt(stability["wind_force_kn"], 3)),
        ("Опрокидывающий момент, кН·м", _fmt(stability["overturning_moment_knm"], 3)),
        ("Удерживающий момент, кН·м", _fmt(stability["resisting_moment_knm"], 3)),
        ("Сила трения, кН", _fmt(stability["sliding_capacity_kn"], 3)),
        ("Среднее давление на основание, кПа", _fmt(stability["avg_foundation_kpa"], 3)),
        ("Максимальное давление на основание, кПа", _fmt(stability["pmax_foundation_kpa"], 3)),
        ("Допускаемое давление, кПа", _fmt(stability["foundation_limit_kpa"], 3)),
    ], title="9. Укрупненная оценка устойчивости и основания")

    check_rows = []
    for item in result["check_items"]:
        value = _fmt(item["value"], 3) if item.get("value") is not None else "—"
        limit = _fmt(item["limit"], 3) if item.get("limit") is not None else "—"
        note = item.get("note") or "—"
        check_rows.append([item["title"], _status_text(item["status"]), value, limit, item.get("unit") or "—", note])
    b.grid_table(
        ["Проверка", "Статус", "Факт", "Предел / ориентир", "Ед.", "Примечание"],
        check_rows,
        title="10. Сводный перечень проверок",
        widths_cm=[4.6, 1.9, 1.8, 2.2, 1.1, 5.0],
        status_col=1,
    )

    flag_rows = []
    for flag in flags:
        flag_rows.append([_status_text(flag["level"]), flag["title"], flag["note"]])
    if flag_rows:
        b.grid_table(
            ["Уровень", "Замечание", "Пояснение"],
            flag_rows,
            title="11. Замечания экспертизы расчетной модели",
            widths_cm=[2.0, 4.8, 9.2],
            status_col=0,
        )
    else:
        b.note_box("Замечания", "Критических и существенных замечаний по введенной расчетной базе не выявлено.", fill=GREEN_LIGHT)

    b.heading("12. Выводы и рекомендации", 1)
    b.bullet(f"Максимальное эксплуатационное напряжение по стенке: {_fmt(summary['sigma_max_mpa'], 3)} МПа.")
    if summary["controlling_belt"] is not None:
        b.bullet(f"Определяющим является пояс № {summary['controlling_belt']}.")
    b.bullet(f"Минимальный коэффициент запаса по номинальной толщине: {_fmt(summary['min_reserve'], 3)}.")
    b.bullet(f"Проверка поясов стенки: {_bool_text(result['checks']['shell_strength_ok'])}.")
    b.bullet(f"Проверка при гидроиспытании: {_bool_text(result['checks']['shell_hydrotest_ok'])}.")
    b.bullet(f"Укрупненная проверка основания: {_bool_text(result['checks']['foundation_ok'])}.")
    final_text = (
        "По введенной модели критических несоответствий не выявлено, однако итоговый документ может быть выпущен как полноценная проектная записка только после закрытия всех отсутствующих исходных данных и выполнения невыполненных нормативных проверок."
        if summary["final_ok"] else
        "По введенной модели выявлены критические или существенные замечания. До выпуска рабочей документации требуется корректировка расчетной базы, типа резервуара и состава проверок."
    )
    b.note_box("Финальный вывод", final_text, fill=GREEN_LIGHT if summary["final_ok"] else YELLOW_LIGHT)

    b.heading("Приложение А. Снимок расчетной модели", 1)
    b.key_value_table([
        ("Автоматически назначенный класс по объему", design_basis.get("reservoir_class_auto") or "—"),
        ("Принятый класс резервуара", design_basis.get("reservoir_class") or "—"),
        ("Тип резервуара", design_basis.get("tank_type") or "—"),
        ("Срок службы, лет", _fmt(design_basis.get("service_life_years"), 0) if design_basis.get("service_life_years") else "—"),
        ("Циклы/год", _fmt(design_basis.get("cycles_per_year"), 0) if design_basis.get("cycles_per_year") else "—"),
        ("Тип местности", design_basis.get("terrain_type") or "—"),
        ("ГО/УЛФ / инертирование", design_basis.get("gas_balance_system") or "—"),
        ("Относительный вакуум, кПа", _fmt(design_basis.get("vacuum_kpa"), 3)),
    ])


def _build_foundation_report(doc: Document, result: dict) -> None:
    b = DocxBuilder(doc)
    b.title_box(
        "ОТЧЕТ ПО УКРУПНЕННОЙ ПРОВЕРКЕ ОСНОВАНИЯ И УСТОЙЧИВОСТИ",
        "Приложение к расчетной модели резервуара",
        [
            ("Объект", result["meta"].get("title") or "РВС"),
            ("Класс резервуара", result["meta"].get("reservoir_class") or "—"),
            ("Дата", datetime.now().strftime("%d.%m.%Y %H:%M:%S")),
        ],
    )
    b.key_value_table([
        ("Среднее давление на основание, кПа", _fmt(result["stability"]["avg_foundation_kpa"], 3)),
        ("Максимальное давление на основание, кПа", _fmt(result["stability"]["pmax_foundation_kpa"], 3)),
        ("Допускаемое давление, кПа", _fmt(result["stability"]["foundation_limit_kpa"], 3)),
        ("Ветровая сила, кН", _fmt(result["stability"]["wind_force_kn"], 3)),
        ("Опрокидывающий момент, кН·м", _fmt(result["stability"]["overturning_moment_knm"], 3)),
        ("Удерживающий момент, кН·м", _fmt(result["stability"]["resisting_moment_knm"], 3)),
        ("Сдвиговая устойчивость, кН", _fmt(result["stability"]["sliding_capacity_kn"], 3)),
        ("Вывод", _bool_text(result["checks"]["foundation_ok"])),
    ], title="Основные расчетные показатели")
    b.note_box(
        "Ограничение",
        "Документ носит характер укрупненной проверки и должен подтверждаться расчетом основания и фундамента по данным инженерно-геологических изысканий и полной схеме сочетаний нагрузок.",
    )


def _build_terms_of_reference(doc: Document, result: dict) -> None:
    b = DocxBuilder(doc)
    geom = result["geometry"]
    inputs = result["inputs"]
    meta = result["meta"]
    bottom = result.get("bottom", {})
    roof = result.get("roof", {})
    masses = result.get("masses", {})
    stability = result.get("stability", {})
    design = result.get("design_basis", {})
    b.title_box(
        "ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ПРОЕКТИРОВАНИЕ РЕЗЕРВУАРА",
        "Сформировано по данным расчетной модели РВС и предназначено для согласования исходных данных",
        [
            ("Наименование", meta.get("title") or "РВС"),
            ("Тип резервуара", meta.get("tank_type") or "РВС"),
            ("Класс резервуара", meta.get("reservoir_class") or "—"),
            ("Дата", datetime.now().strftime("%d.%m.%Y %H:%M:%S")),
        ],
    )
    b.note_box(
        "Назначение документа",
        "Документ является предварительным техническим заданием/опросным листом для вертикального стального резервуара. Перед передачей в проектирование его необходимо дополнить реквизитами заказчика, схемой площадки, требованиями пожарной и промышленной безопасности, КИПиА, обвязкой и комплектом рабочей документации.",
        fill=ACCENT_LIGHT,
    )
    b.key_value_table([
        ("Тип резервуара", meta.get("tank_type") or "РВС"),
        ("Класс резервуара", meta.get("reservoir_class") or "—"),
        ("Расчетный срок службы, лет", _fmt(inputs.get("service_life_years"), 0) if inputs.get("service_life_years") else "—"),
        ("Оборачиваемость, циклов/год", _fmt(inputs.get("cycles_per_year"), 0) if inputs.get("cycles_per_year") else "—"),
        ("Нормативная база", "ГОСТ 31385-2023; СП 20.13330; СП 16.13330; СП 22.13330; СП 14.13330"),
    ], title="1. Общие данные")
    b.key_value_table([
        ("Хранимый продукт", str(inputs["medium"])),
        ("Плотность продукта, кг/м³", _fmt(inputs["density_kg_m3"], 1)),
        ("Рабочий уровень налива, мм", _fmt(inputs["fill_level_mm"], 0)),
        ("Заполнение, %", _fmt(inputs.get("fill_percent"), 1)),
        ("Избыточное давление, МПа", _fmt(inputs["p_gas_mpa"], 5)),
        ("Давление гидроиспытания, МПа", _fmt(inputs.get("p_test_mpa"), 5)),
        ("Относительный вакуум, кПа", _fmt(inputs.get("vacuum_kpa"), 3)),
        ("ГО/УЛФ / инертирование", str(inputs.get("gas_balance_system") or "—")),
        ("Минимальная температура продукта, °C", _fmt_any(inputs.get("product_temp_min_c"), 0)),
        ("Максимальная температура продукта, °C", _fmt_any(inputs.get("product_temp_max_c"), 0)),
    ], title="2. Среда и режим работы")
    b.key_value_table([
        ("Ветровое давление w0, кПа", _fmt(inputs["wind_kpa"], 3)),
        ("Снеговая нагрузка Sg, кПа", _fmt(inputs["snow_kpa"], 3)),
        ("Сейсмичность", str(inputs["seismic"])),
        ("Тип местности", str(meta.get("terrain_type") or design.get("terrain_type") or "—")),
        ("Расчетное сопротивление основания, кПа", _fmt(inputs.get("foundation_limit_kpa"), 1)),
        ("Коэффициент трения по основанию", _fmt(inputs.get("friction_coeff"), 2)),
        ("Максимальная температура наружного воздуха, °C", _fmt_any(inputs.get("ambient_temp_max_c"), 0)),
    ], title="3. Район установки и воздействия")
    b.key_value_table([
        ("Внутренний диаметр стенки, мм", _fmt(inputs["diameter_mm"], 0)),
        ("Высота стенки, мм", _fmt(inputs["height_mm"], 0)),
        ("Номинальный объем, м³", _fmt(geom["full_volume_m3"], 3)),
        ("Рабочий объем, м³", _fmt(geom["useful_volume_m3"], 3)),
        ("Площадь днища в плане, м²", _fmt(geom.get("plan_area_m2"), 3)),
        ("Длина окружности стенки, м", _fmt(geom.get("circumference_m"), 3)),
    ], title="4. Основные размеры")

    belt_rows = []
    for item in result.get("belts", []):
        belt_rows.append([
            _fmt(item.get("belt"), 0),
            _fmt(item.get("height_mm"), 0),
            _fmt(item.get("thickness_nominal_mm"), 1),
            _fmt(item.get("thickness_corrosion_mm"), 1),
            _fmt(item.get("thickness_minus_mm"), 1),
            _fmt(item.get("thickness_required_nominal_mm"), 2),
            _fmt(item.get("stress_service_mpa"), 2),
            _fmt(item.get("reserve_ratio"), 2),
        ])
    b.grid_table(
        ["Пояс", "Высота, мм", "t принятая, мм", "Корр., мм", "Минус, мм", "t треб., мм", "Напр., МПа", "Запас"],
        belt_rows or [["—", "—", "—", "—", "—", "—", "—", "—"]],
        title="5. Стенка резервуара",
        widths_cm=[1.2, 2.0, 2.0, 1.7, 1.7, 2.0, 2.0, 1.5],
    )
    b.key_value_table([
        ("Исполнение днища", str(bottom.get("execution") or "—")),
        ("Толщина центральной части, мм", _fmt(bottom.get("bottom_t_mm"), 1)),
        ("Толщина окрайки, мм", _fmt(bottom.get("ring_t_mm"), 1)),
        ("Диаметр днища, мм", _fmt(bottom.get("bottom_diameter_mm"), 0)),
        ("Ширина окрайки, мм", _fmt(bottom.get("ring_width_mm"), 0)),
        ("Диаметр центральной части, мм", _fmt(bottom.get("center_diameter_mm"), 0)),
        ("Масса центральной части, кг", _fmt(bottom.get("center_mass_kg"), 1)),
        ("Масса окрайки, кг", _fmt(bottom.get("ring_mass_kg"), 1)),
        ("Прибавка на коррозию днища, мм", _fmt(bottom.get("bottom_corr_mm"), 1)),
    ], title="6. Днище")
    b.key_value_table([
        ("Тип крыши", str(roof.get("type") or "—")),
        ("Уклон / угол, град", _fmt(roof.get("angle_deg"), 2)),
        ("Толщина настила, мм", _fmt(roof.get("deck_t_mm"), 1)),
        ("Прибавка на коррозию крыши, мм", _fmt(roof.get("deck_corr_mm"), 1)),
        ("Расчетная масса крыши, кг", _fmt(roof.get("mass_kg"), 1)),
    ], title="7. Крыша")
    b.key_value_table([
        ("Масса стенки, кг", _fmt(masses.get("shell_kg"), 1)),
        ("Масса днища, кг", _fmt(masses.get("bottom_kg"), 1)),
        ("Масса крыши, кг", _fmt(masses.get("roof_kg"), 1)),
        ("Металлоконструкции и оборудование, кг", _fmt(masses.get("metal_kg"), 1)),
        ("Теплоизоляция, кг", _fmt(masses.get("insulation_kg"), 1)),
        ("Масса продукта, кг", _fmt(masses.get("product_kg"), 1)),
        ("Снеговая масса, кг", _fmt(masses.get("snow_kg"), 1)),
        ("Полная расчетная масса, кг", _fmt(masses.get("total_kg"), 1)),
        ("Вертикальная нагрузка, кН", _fmt(masses.get("total_load_kn"), 2)),
    ], title="8. Массы и нагрузки")
    b.key_value_table([
        ("Ветровая сила, кН", _fmt(stability.get("wind_force_kn"), 3)),
        ("Опрокидывающий момент, кН·м", _fmt(stability.get("overturning_moment_knm"), 3)),
        ("Удерживающий момент, кН·м", _fmt(stability.get("resisting_moment_knm"), 3)),
        ("Сдвиговая несущая способность, кН", _fmt(stability.get("sliding_capacity_kn"), 3)),
        ("Среднее давление на основание, кПа", _fmt(stability.get("avg_foundation_kpa"), 3)),
        ("Максимальное давление на основание, кПа", _fmt(stability.get("pmax_foundation_kpa"), 3)),
    ], title="9. Предварительные проверки")
    check_rows = []
    for item in result.get("check_items", []):
        check_rows.append([
            str(item.get("title") or item.get("code") or "—"),
            _status_text(str(item.get("status") or "")),
            _fmt_any(item.get("value"), 3),
            _fmt_any(item.get("limit"), 3),
            str(item.get("unit") or "—"),
            str(item.get("note") or "—"),
        ])
    b.grid_table(
        ["Проверка", "Статус", "Значение", "Предел", "Ед.", "Комментарий"],
        check_rows or [["—", "—", "—", "—", "—", "—"]],
        title="10. Контрольные условия",
        widths_cm=[4.4, 1.8, 1.8, 1.8, 1.0, 5.4],
        status_col=1,
    )
    b.heading("11. Требования к комплектации, документации и контролю", 1)
    for item in [
        "В составе проектной/рабочей документации предусмотреть общие виды, узлы стенки, днища, крыши, лестниц, площадок, патрубков, люков и заземления.",
        "Схему патрубков, люков, дыхательной арматуры, пеногенераторов, КИПиА и пожарного оборудования согласовать с технологической частью проекта.",
        "Назначить требования к материалам, сварке, контролю сварных соединений, гидроиспытанию, антикоррозионной защите, маркировке и консервации.",
        "Для основания и фундамента выполнить отдельную проверку по материалам инженерно-геологических изысканий и расчетной схеме площадки.",
        "Уточнить требования промышленной, пожарной и экологической безопасности для конкретного продукта и объекта.",
    ]:
        b.bullet(item)
    b.note_box(
        "Основание формы",
        "Структура ТЗ ориентирована на исходные данные и требования ГОСТ 31385-2023 для вертикальных стальных резервуаров и дополнена расчетными показателями калькулятора.",
    )


def _extract_rgs_payload(payload: dict) -> dict:
    if isinstance(payload.get("payload"), dict):
        return payload["payload"]
    if isinstance(payload.get("model"), dict):
        return payload["model"]
    return payload


def _rgs_type_label(value: str) -> str:
    return {
        "rgsn": "РГСН, наземный",
        "rgsp": "РГСП, подземный",
        "rgsnd": "РГСНД, наземный двустенный",
        "rgspd": "РГСПД, подземный двустенный",
    }.get(str(value or "").lower(), str(value or "РГС"))


def _rgs_head_label(value: str) -> str:
    return {
        "flat": "плоские",
        "cone": "конические",
        "truncated_cone": "усеченно-конические",
    }.get(str(value or ""), str(value or "—"))


def _rgs_status(value: str) -> str:
    normalized = str(value or "").upper()
    if "PASS" in normalized or normalized == "OK":
        return "OK"
    if "FAIL" in normalized:
        return "НЕ ОК"
    if "WARN" in normalized:
        return "Внимание"
    return str(value or "—")


def _build_rgs_calculation_report(doc: Document, result: dict, payload: dict) -> None:
    b = DocxBuilder(doc)
    summary = result.get("summary", {})
    details = result.get("details", {})
    geometry = details.get("geometry", {})
    shell = details.get("shell", {})
    head = details.get("head", {})
    supports = details.get("supports", {})
    pressures = details.get("pressures", {})
    masses = details.get("masses", {})

    b.title_box(
        "РАСЧЕТ РЕЗЕРВУАРА РГС",
        "Автоматически сформированный расчет по данным калькулятора",
        [
            ("Тип резервуара", _rgs_type_label(payload.get("tank_type"))),
            ("Продукт", str(payload.get("product_name") or "—")),
            ("Геометрия", f"D = {_fmt(payload.get('D', 0) * 1000, 0)} мм; L = {_fmt(payload.get('total_length_m', 0) * 1000, 0)} мм"),
            ("Дата формирования", datetime.now().strftime("%d.%m.%Y %H:%M:%S")),
        ],
    )
    b.note_box("Статус", "Документ формируется как инженерный протокол калькулятора РГС. Для выпуска рабочей документации требуется проверка проектировщиком и дополнение недостающими разделами проекта.", fill=ACCENT_LIGHT)
    b.key_value_table([
        ("Геометрический объем, м3", _fmt(summary.get("full_volume_m3"), 3)),
        ("Объем продукта, м3", _fmt(summary.get("product_volume_m3"), 3)),
        ("Высота заполнения, мм", _fmt((summary.get("fill_height_m") or 0) * 1000.0, 0)),
        ("Сухая масса, кг", _fmt(summary.get("dry_mass_kg"), 1)),
        ("Рабочая масса, кг", _fmt(summary.get("operating_mass_kg"), 1)),
        ("Масса при гидроиспытании, кг", _fmt(summary.get("hydrotest_mass_kg"), 1)),
    ], title="1. Сводные показатели")
    b.key_value_table([
        ("Внутренний диаметр D, мм", _fmt(payload.get("D", 0) * 1000.0, 0)),
        ("Общая длина, мм", _fmt(payload.get("total_length_m", 0) * 1000.0, 0)),
        ("Длина обечайки, мм", _fmt((geometry.get("shell_length_m") or 0) * 1000.0, 0)),
        ("Тип днищ", _rgs_head_label(payload.get("head_type"))),
        ("Расчетная глубина днища, мм", _fmt((geometry.get("head_projection_m") or payload.get("head_projection_m") or 0) * 1000.0, 0)),
        ("Материал", str(payload.get("material") or "—")),
    ], title="2. Геометрия и материал")
    b.key_value_table([
        ("Обечайка, выбранная t, мм", _fmt(shell.get("nominal_mm"), 1)),
        ("Обечайка, эффективная t, мм", _fmt(shell.get("effective_mm"), 2)),
        ("Обечайка, прибавка на коррозию, мм", _fmt(shell.get("corrosion_allowance_mm"), 1)),
        ("Днище, выбранная t, мм", _fmt(head.get("nominal_mm"), 1)),
        ("Днище, прибавка на коррозию, мм", _fmt(head.get("corrosion_allowance_mm"), 1)),
        ("Минусовой допуск проката, мм", _fmt(payload.get("minus_tolerance_mm"), 1)),
        ("Кольца жесткости / диафрагмы, шт.", _fmt(shell.get("ring_count"), 0)),
    ], title="3. Толщины и элементы жесткости")
    b.key_value_table([
        ("Внутреннее давление с гидростатикой, МПа", _fmt(pressures.get("internal_total_mpa"), 4)),
        ("Вакуум / внешнее давление, МПа", _fmt(pressures.get("vacuum_mpa"), 4)),
        ("Требуемое внешнее давление обечайки, МПа", _fmt(pressures.get("external_required_mpa"), 4)),
        ("Требуемое внешнее давление днища, МПа", _fmt(pressures.get("head_external_required_mpa"), 4)),
        ("Вертикальное давление грунта, кПа", _fmt(pressures.get("soil_vertical_kpa"), 2)),
        ("Боковое давление грунта, кПа", _fmt(pressures.get("soil_horizontal_kpa"), 2)),
    ], title="4. Расчетные нагрузки")
    b.key_value_table([
        ("Обечайка, кг", _fmt(masses.get("shell_mass_kg"), 1)),
        ("Днища, кг", _fmt(masses.get("head_total_mass_kg"), 1)),
        ("Патрубки, кг", _fmt(masses.get("nozzle_mass_kg"), 1)),
        ("Опоры, кг", _fmt(masses.get("supports_mass_kg"), 1)),
        ("Крепление теплоизоляции, кг", _fmt(masses.get("insulation_supports_mass_kg"), 1)),
        ("Металл всего, кг", _fmt(summary.get("steel_mass_kg"), 1)),
    ], title="5. Массы")
    b.key_value_table([
        ("Опоры, принято, шт.", _fmt(supports.get("support_count"), 0)),
        ("Опоры, рекомендовано, шт.", _fmt(supports.get("recommended_support_count"), 0)),
        ("Высота седла, мм", _fmt((supports.get("saddle_height_m") or 0) * 1000.0, 0)),
        ("Масса одной опоры, кг", _fmt(supports.get("actual_weight_each_kg"), 1)),
        ("Давление под седлом, кПа", _fmt(supports.get("foundation_pressure_kpa"), 2)),
        ("Дуга опирания, град", _fmt(supports.get("contact_angle_deg"), 0)),
    ], title="6. Опоры")
    check_rows = []
    for item in result.get("checks", []):
        if item.get("code") == "RGS_SUPPORT_BEARING":
            continue
        check_rows.append([str(item.get("title") or item.get("code") or "—"), _rgs_status(item.get("result")), _fmt_any(item.get("value"), 4), _fmt_any(item.get("limit"), 4), str(item.get("unit") or "—"), str(item.get("reference") or "—")])
    b.grid_table(["Проверка", "Статус", "Результат", "Допустимо / требуется", "Ед.", "Нормативная ссылка"], check_rows or [["Проверки", "—", "—", "—", "—", "—"]], title="7. Проверки", widths_cm=[4.7, 1.8, 2.1, 2.3, 1.0, 4.4], status_col=1)
    protocol_rows = []
    for item in result.get("protocol", []):
        if item.get("code") == "RGS_SUPPORT_BEARING":
            continue
        protocol_rows.append([str(item.get("title") or item.get("code") or "—"), str(item.get("formula") or "—"), _fmt_any(item.get("value"), 4), _fmt_any(item.get("margin"), 3), str(item.get("note") or "—")])
    b.grid_table(["Раздел", "Формула / методика", "Результат", "Запас", "Пояснение"], protocol_rows or [["—", "—", "—", "—", "—"]], title="8. Протокол расчета", widths_cm=[3.5, 4.4, 1.8, 1.4, 5.2])
    b.heading("9. Нормативная база", 1)
    for item in result.get("normative", []):
        b.bullet(str(item))


def _build_rgs_terms_of_reference(doc: Document, result: dict, payload: dict) -> None:
    b = DocxBuilder(doc)
    summary = result.get("summary", {})
    details = result.get("details", {})
    supports = details.get("supports", {})
    insulation = details.get("insulation_supports", {})
    nozzles = details.get("nozzles", {})
    flotation = details.get("flotation", {})
    b.title_box(
        "ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА РЕЗЕРВУАР РГС",
        "Форма подготовлена на основе опросного листа ГОСТ 17032-2022, приложение А",
        [("Наименование", _rgs_type_label(payload.get("tank_type"))), ("Назначение", "хранение жидкого продукта"), ("Продукт", str(payload.get("product_name") or "—")), ("Дата формирования", datetime.now().strftime("%d.%m.%Y %H:%M:%S"))],
    )
    b.note_box("Назначение документа", "Это предварительное техническое задание/опросный лист для согласования исходных данных. Перед передачей в производство его нужно дополнить реквизитами заказчика, требованиями площадки, КИПиА, пожарной безопасностью, схемой патрубков и комплектом рабочей документации.", fill=ACCENT_LIGHT)
    b.key_value_table([
        ("Тип резервуара", _rgs_type_label(payload.get("tank_type"))),
        ("Исполнение", "двустенное" if "d" in str(payload.get("tank_type", "")) else "одностенное"),
        ("Установка", "подземная" if str(payload.get("tank_type", "")).startswith("rgsp") else "наземная"),
        ("Регион / город", ", ".join([str(payload.get("site_region") or ""), str(payload.get("site_city") or "")]).strip(", ") or "—"),
        ("Нормативная база", "ГОСТ 17032-2022; ГОСТ 34347; ГОСТ 34233.1/2; СП 20.13330; СП 14.13330"),
    ], title="1. Общие данные")
    b.key_value_table([
        ("Хранимый продукт", str(payload.get("product_name") or "—")),
        ("Плотность продукта, кг/м3", _fmt(payload.get("rho"), 1)),
        ("Температура продукта, °C", _fmt(payload.get("temperature_c"), 0)),
        ("Заполнение", f"{_fmt_any(payload.get('fill_value'), 1)} {'м' if payload.get('fill_mode') == 'level' else '%'}"),
        ("Рабочее внутреннее давление, МПа", _fmt(payload.get("design_pressure_mpa"), 4)),
        ("Вакуум, МПа", _fmt(payload.get("vacuum_mpa"), 4)),
    ], title="2. Среда и режим работы")
    b.key_value_table([
        ("Ветровой район / w0, кПа", f"{payload.get('wind_region') or '—'} / {_fmt(payload.get('w0_kpa'), 2)}"),
        ("Снеговой район / Sg, кПа", f"{payload.get('snow_region') or '—'} / {_fmt(payload.get('sg_kpa'), 2)}"),
        ("Минус пятидневки, °C", _fmt(payload.get("t5_c"), 0)),
        ("Абсолютный минимум, °C", _fmt(payload.get("tmin_abs_c"), 0)),
        ("Сейсмичность", str(payload.get("seismic") or "—")),
        ("ОСР-97", str(payload.get("seis_level") or "—")),
    ], title="3. Район установки")
    b.key_value_table([
        ("Геометрический объем, м3", _fmt(summary.get("full_volume_m3"), 3)),
        ("Внутренний диаметр D, мм", _fmt(payload.get("D", 0) * 1000.0, 0)),
        ("Общая длина, мм", _fmt(payload.get("total_length_m", 0) * 1000.0, 0)),
        ("Тип днищ", _rgs_head_label(payload.get("head_type"))),
        ("Малое основание усеченного конуса, мм", _fmt(payload.get("head_small_diameter_m", 0) * 1000.0, 0)),
        ("Материал", str(payload.get("material") or "—")),
    ], title="4. Основные размеры")
    b.key_value_table([
        ("Толщина обечайки, мм", _fmt(payload.get("shell_nominal_mm"), 1)),
        ("Прибавка на коррозию обечайки, мм", _fmt(payload.get("corr_mm"), 1)),
        ("Толщина днища, мм", _fmt(payload.get("head_nominal_mm"), 1)),
        ("Прибавка на коррозию днища, мм", _fmt(payload.get("head_allowances_mm"), 1)),
        ("Минусовой допуск проката, мм", _fmt(payload.get("minus_tolerance_mm"), 1)),
        ("Кольца жесткости / диафрагмы, шт.", _fmt(payload.get("ring_count"), 0)),
    ], title="5. Материал и толщины")
    nozzle_rows = []
    for item in payload.get("nozzles") or []:
        if int(item.get("count") or 0) <= 0:
            continue
        nozzle_rows.append([str(item.get("name") or "Патрубок"), _fmt(item.get("count"), 0), _fmt(item.get("dn_mm"), 0), _fmt(item.get("length_mm"), 0), _fmt(item.get("thickness_mm"), 1)])
    b.grid_table(["Назначение", "Кол-во", "DN, мм", "Длина, мм", "t, мм"], nozzle_rows or [["Патрубки", _fmt(nozzles.get("count"), 0), "—", "—", "—"]], title="6. Патрубки", widths_cm=[5.8, 1.7, 2.0, 2.5, 2.0])
    b.key_value_table([
        ("Опоры, принято / рекомендовано, шт.", f"{_fmt(supports.get('support_count'), 0)} / {_fmt(supports.get('recommended_support_count'), 0)}"),
        ("Высота опоры, мм", _fmt((supports.get("saddle_height_m") or 0) * 1000.0, 0)),
        ("Ширина седла, мм", _fmt(payload.get("saddle_width_m", 0) * 1000.0, 0)),
        ("Масса одной опоры, кг", _fmt(supports.get("actual_weight_each_kg"), 1)),
        ("Лестница внутренняя", _fmt_any(payload.get("ladder"))),
        ("Площадка лестницы", _fmt_any(payload.get("platform"))),
    ], title="7. Опоры и навесное оборудование")
    b.key_value_table([
        ("Теплоизоляция", _fmt_any(payload.get("insulation_enabled"))),
        ("Тип изоляции", "ППУ без металлического крепления" if str(payload.get("tank_type", "")).startswith("rgsp") else "теплоизоляция с креплением"),
        ("Толщина теплоизоляции, мм", _fmt(payload.get("insulation_thickness_mm"), 0)),
        ("Плотность теплоизоляции, кг/м3", _fmt(payload.get("insulation_density_kg_m3"), 0)),
        ("Кольца крепления ТИ, шт.", _fmt(insulation.get("shell_ring_count"), 0)),
        ("Т-образные проставки по обечайке, шт.", _fmt(insulation.get("stand_count"), 0)),
        ("Покрытие, мм", _fmt(payload.get("coating_thickness_mm"), 1)),
    ], title="8. Теплоизоляция и покрытия")
    if str(payload.get("tank_type", "")).startswith("rgsp"):
        b.key_value_table([
            ("Тип грунта", str(payload.get("soil_preset") or "—")),
            ("Глубина до верхней образующей, мм", _fmt(payload.get("burial_depth_top_m", 0) * 1000.0, 0)),
            ("Плотность грунта, кг/м3", _fmt(payload.get("soil_density_kg_m3"), 0)),
            ("Угол внутреннего трения, град", _fmt(payload.get("soil_phi_deg"), 0)),
            ("Грунтовые воды", _fmt_any(payload.get("groundwater"))),
            ("Уровень грунтовых вод, мм", _fmt(payload.get("groundwater_level_m", 0) * 1000.0, 0)),
            ("Погружение в воду, мм", _fmt(flotation.get("submerged_height_m", 0) * 1000.0, 0)),
            ("Хомуты от всплытия, шт.", _fmt(flotation.get("strap_count"), 0)),
        ], title="9. Подземная установка")
    b.heading("10. Требования к документации и контролю", 1)
    for item in ["Разработать КМ/КМД или рабочую документацию в составе, согласованном с заказчиком.", "Указать схему расположения патрубков, люков, опор, строповочных элементов и заземления.", "Назначить требования к сварке, НК, гидроиспытанию, консервации и маркировке.", "Уточнить требования промышленной, пожарной и экологической безопасности для конкретного продукта."]:
        b.bullet(item)
    b.note_box("Основание формы", "Структура ТЗ ориентирована на опросный лист для горизонтальных цилиндрических резервуаров из приложения А ГОСТ 17032-2022 и дополнена расчетными данными калькулятора.")


def create_strength_report(payload: dict) -> Path:
    result = calculate_strength(payload)
    out_dir = Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / _safe_filename("rvs_strength_report")
    doc = Document()
    _build_strength_report(doc, result)
    doc.save(path)
    return path


def create_foundation_report(payload: dict) -> Path:
    result = calculate_strength(payload)
    out_dir = Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / _safe_filename("foundation_report")
    doc = Document()
    _build_foundation_report(doc, result)
    doc.save(path)
    return path


def create_terms_of_reference(payload: dict) -> Path:
    result = calculate_strength(payload)
    out_dir = Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / _safe_filename("terms_of_reference")
    doc = Document()
    _build_terms_of_reference(doc, result)
    doc.save(path)
    return path


def create_rgs_calculation_report(payload: dict) -> Path:
    rgs_payload = _extract_rgs_payload(payload)
    result = calculate_rgs(rgs_payload)
    out_dir = Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / _safe_filename("rgs_calculation_report")
    doc = Document()
    _build_rgs_calculation_report(doc, result, rgs_payload)
    doc.save(path)
    return path


def create_rgs_terms_of_reference(payload: dict) -> Path:
    rgs_payload = _extract_rgs_payload(payload)
    result = calculate_rgs(rgs_payload)
    out_dir = Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / _safe_filename("rgs_terms_of_reference")
    doc = Document()
    _build_rgs_terms_of_reference(doc, result, rgs_payload)
    doc.save(path)
    return path
